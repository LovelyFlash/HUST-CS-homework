from docx import Document
from docx.enum.text import WD_COLOR_INDEX

document = Document()

# 新闻内容部分
document.add_paragraph("In contrast to his first term in office, President Trump has faced relatively little opposition since he was sworn in 76 days ago. Republicans, many Democrats, tech giants, universities, and even law firms, have been accused of caving into his demands. The only real challenge has come from a handful of judges. But on Saturday, tens of thousands of people took to the streets across America to voice their anger at the way the country is heading. The so-called hands-off events were organized in every congressional district with protesters criticizing the doge cuts, Donald Trump's economic policy and what some see as a slide towards authoritarianism. These protesters are among those concerned that America's democracy is at risk.")
document.add_paragraph("【Q1】")
document.add_paragraph()

document.add_paragraph("\"Stock market's crashing. The economy is going to crash. It's already crashing.\"")
document.add_paragraph()
document.add_paragraph("\"And it's all about Trump, his actions, his stupidity, his mistakes.\"")
document.add_paragraph()
document.add_paragraph("\"One of my major concerns is how much disinformation that they are perpetuating out in the public, that basic science and basic facts that our democracy depends on is being torn down. And when that happens, people get hurt.\"")
document.add_paragraph()
document.add_paragraph("Well, for his part, President Trump took to social media to call on Americans to hang tough in the face of the turmoil caused by his sweeping import tariffs. And he still has plenty of support from the likes of Brian Pana, Becca, a retired auto worker from Michigan.")
document.add_paragraph()
document.add_paragraph("\"We've endured this pain in the United States for 30 or 40 years of seeing our factories close. So we can certainly endure two or three years of economic pain while the supply chains are reorganized and the companies move their production lines back to the United States. I couldn't give a rat's rear end about the stock markets. I personally believe they were overvalued and correction was in order.\"")
document.add_paragraph()
document.add_paragraph("Our North America correspondent Peter Bows followed the day's developments from Washington. These protests were widespread. More than a thousand individual demonstrations across the country. That is significant.")
document.add_paragraph("【Q2】")
document.add_paragraph()

document.add_paragraph("The biggest single day of protests since Donald Trump returned to the White House, perhaps surprising to some that it's taken so long given this country is still so polarized in terms of the nature of politics here so long that it's taken for large groups of people to get together and take to the streets to express their anger in this way.")
document.add_paragraph()
document.add_paragraph("But they were protesting about many different policies, the executive orders, the widespread sacking of federal workers, the breaking up of the department of education and more. And I think today was simply an opportunity to vent about all of that as well as the president's most recently and of course, arguably biggest bombshells since he took office. And that is the sweeping tariffs that have been brought in on the import of goods into the US.")
document.add_paragraph()

# 听力题部分
document.add_paragraph("### Listening Questions:")
p = document.add_paragraph()
run = p.add_run("Q1. What were the main concerns of the protesters during the nationwide demonstrations?")
run.bold = True
document.add_paragraph("   A. Climate change and gun control")
document.add_paragraph("   B. Economic policies, disinformation, and threats to democracy")
p_answer1 = document.add_paragraph()
run_answer1 = p_answer1.add_run("      【答案依据】：Economic policies, disinformation, and threats to democracy")
run_answer1.font.highlight_color = WD_COLOR_INDEX.YELLOW
document.add_paragraph("   C. Foreign relations and war")
document.add_paragraph("   D. Healthcare and education")
document.add_paragraph()

p = document.add_paragraph()
run = p.add_run("Q2. What did retired auto worker Brian Pana suggest about the economic situation?")
run.bold = True
document.add_paragraph("   A. The economy cannot recover without foreign help")
document.add_paragraph("   B. The stock market should be prioritized")
document.add_paragraph("   C. Short-term pain is acceptable for long-term industrial recovery")
p_answer2 = document.add_paragraph()
run_answer2 = p_answer2.add_run("      【答案依据】：Short-term pain is acceptable for long-term industrial recovery")
run_answer2.font.highlight_color = WD_COLOR_INDEX.YELLOW
document.add_paragraph("   D. Tariffs should be immediately canceled")
document.add_paragraph()

# 生词表部分
document.add_paragraph("### Vocabulary:")
table = document.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Word/Phrase"
hdr_cells[1].text = "Definition"
row = table.add_row().cells
row[0].text = "sweeping tariffs"
row[1].text = "全面征收的关税"
row = table.add_row().cells
row[0].text = "caving into"
row[1].text = "屈服于"
row = table.add_row().cells
row[0].text = "authoritarianism"
row[1].text = "专制主义"
row = table.add_row().cells
row[0].text = "disinformation"
row[1].text = "虚假信息，假消息"
row = table.add_row().cells
row[0].text = "perpetuate"
row[1].text = "使持续，使长存（尤指不好的事）"
row = table.add_row().cells
row[0].text = "polarized"
row[1].text = "两极分化的"
row = table.add_row().cells
row[0].text = "vent (about sth)"
row[1].text = "发泄（情绪）"
row = table.add_row().cells
row[0].text = "sacking"
row[1].text = "解雇"
row = table.add_row().cells
row[0].text = "bombshell"
row[1].text = "爆炸性新闻"

document.save("news_report_updated.docx")
